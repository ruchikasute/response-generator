# import streamlit as st
# import os
# import time
# import re
# from io import BytesIO
# from dotenv import load_dotenv
# from PyPDF2 import PdfReader
# import docx
# from docx import Document
# from openai import AzureOpenAI
# from langchain_openai import AzureOpenAIEmbeddings
# from langchain_chroma import Chroma
# from langchain_core.documents import Document as LDocument
# from docx.shared import Inches, RGBColor
# from docx.oxml import OxmlElement
# from docx.oxml.ns import qn
# from docx.shared import Pt
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.enum.table import WD_ALIGN_VERTICAL
# from Modules.prompts import (
#     get_executive_summary_and_objective_prompt,
#     get_scope_prereq_assumptions_prompt,     get_resource_schedule_and_commercial_prompt, get_communication_plan_prompt
# )


# # -------------------------------------------------------
# # 1. SETUP
# # -------------------------------------------------------
# load_dotenv()
# KNOWLEDGE_FOLDER = "Knowledge_Repo"
# PERSIST_DIR = "chroma_db"

# st.set_page_config(page_title="📄 RFP Executive Summary Generator", layout="wide")
# st.markdown("<h1 style='text-align:center; color:#4B0082;'>📄 Executive Summary Generator</h1>", unsafe_allow_html=True)
# st.markdown("---")

# # -------------------------------------------------------
# # 2. UTILITIES
# # -------------------------------------------------------

# def extract_text(file):
#     """Extract text from PDF or DOCX"""
#     if file.name.endswith(".pdf"):
#         reader = PdfReader(file)
#         return "\n".join([p.extract_text() or "" for p in reader.pages])
#     elif file.name.endswith(".docx"):
#         doc = docx.Document(file)
#         return "\n".join([p.text for p in doc.paragraphs])
#     return ""

# def build_knowledge_base(folder=KNOWLEDGE_FOLDER, persist_dir=PERSIST_DIR):
#     """Build or load persistent Chroma vector DB from Knowledge_Repo"""
#     os.makedirs(folder, exist_ok=True)
#     os.makedirs(persist_dir, exist_ok=True)

#     embedding_model = AzureOpenAIEmbeddings(
#         model="text-embedding-ada-002",
#         azure_endpoint=os.getenv("AZURE_OPENAI_EMD_ENDPOINT"),
#         api_key=os.getenv("AZURE_OPENAI_EMD_KEY"),
#         api_version=os.getenv("AZURE_OPENAI_EMD_VERSION")
#     )

#     if os.listdir(persist_dir):
#         return Chroma(
#             embedding_function=embedding_model,
#             persist_directory=persist_dir,
#             collection_name="rfp_responses"
#         )

#     docs = []
#     for f in os.listdir(folder):
#         if f.endswith((".pdf", ".docx")):
#             path = os.path.join(folder, f)
#             text = extract_text(open(path, "rb"))
#             if text.strip():
#                 docs.append(LDocument(page_content=text, metadata={"source": f}))

#     if not docs:
#         raise ValueError(f"No readable files found in {folder}")

#     return Chroma.from_documents(
#         documents=docs,
#         embedding=embedding_model,
#         persist_directory=persist_dir,
#         collection_name="rfp_responses"
#     )



# def insert_executive_summary_into_template(
#     template_path,
#     summary_text,
#     objective_text=None,
#     scope_text=None,
#     resource_schedule_text=None,
#     communication_plan_text=None,
# ):
#     """
#     Replace placeholders in the template:
#     <<EXEC_SUMMARY>>, <<OBJECTIVE>>, <<SCOPE_TEXT>>, <<RESOURCE_SCHEDULE>>, <<COMMUNICATION_PLAN>>
#     Supports formatted tables with styled headers (blue) and light gray rows.
#     """

#     doc = Document(template_path)

#     def set_cell_shading(cell, fill_color):
#         """Add shading (background color) to a table cell."""
#         tc_pr = cell._element.tcPr
#         shd = OxmlElement("w:shd")
#         shd.set(qn("w:val"), "clear")
#         shd.set(qn("w:color"), "auto")
#         shd.set(qn("w:fill"), fill_color)
#         tc_pr.append(shd)

#     def set_table_border_white(table, cell_margin=150):
#         """Set all table borders to white (for clean, minimal look)."""
#         tbl = table._element
#         tbl_pr = tbl.tblPr
#         tbl_borders = OxmlElement("w:tblBorders")

#         for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
#             border_el = OxmlElement(f"w:{border_name}")
#             border_el.set(qn("w:val"), "single")
#             border_el.set(qn("w:sz"), "4")  # thin border
#             border_el.set(qn("w:space"), "0")
#             border_el.set(qn("w:color"), "FFFFFF")  # white
#             tbl_borders.append(border_el)

#         tbl_pr.append(tbl_borders)


#     def insert_styled_table(parent, headers, rows):
#         """Create a table styled similar to RFP objective section."""
#         table = parent.add_table(rows=len(rows) + 1, cols=len(headers))
#         table.style = "Table Grid"
#         table.autofit = True

#         # Header row styling
#         hdr_cells = table.rows[0].cells
#         for i, h in enumerate(headers):
#             hdr_cells[i].text = h.strip()
#             set_cell_shading(hdr_cells[i], "008FD3")  # blue header
#             for run in hdr_cells[i].paragraphs[0].runs:
#                 run.font.bold = True
#                 run.font.color.rgb = RGBColor(255, 255, 255)
#             hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
#             hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

#         # Data rows
#         for r, row_data in enumerate(rows):
#             cells = table.rows[r + 1].cells
#             for c, val in enumerate(row_data):
#                 cells[c].text = str(val).strip()
#                 set_cell_shading(cells[c], "E7EEF7")  # light gray row
#                 cells[c].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
#                 cells[c].paragraphs[0].alignment = (
#                     WD_ALIGN_PARAGRAPH.LEFT if c == 0 else WD_ALIGN_PARAGRAPH.LEFT
#                 )

#         # Set uniform width
#         for row in table.rows:
#             for cell in row.cells:
#                 cell.width = Inches(3)

#         # Apply white borders
#         set_table_border_white(table)

#         return table


#     def replace_placeholder(doc, placeholder, new_text):
#         if not new_text:
#             return

#         for para in doc.paragraphs:
#             if placeholder in "".join(run.text for run in para.runs):
#                 parent = para._element.getparent()
#                 idx = parent.index(para._element)
#                 parent.remove(para._element)

#                 lines = [line.strip() for line in new_text.split("\n") if line.strip()]
#                 new_elements = []  # collect to insert once

#                 i = 0
#                 while i < len(lines):
#                     line = lines[i]

#                     # Markdown-style table
#                     if line.startswith("|") and "|" in line:
#                         table_lines = []
#                         while i < len(lines) and lines[i].startswith("|"):
#                             table_lines.append(lines[i])
#                             i += 1
#                         headers = [h.strip("* ") for h in table_lines[0].strip("|").split("|")]
#                         rows = [
#                             [c.strip() for c in r.strip("|").split("|")]
#                             for r in table_lines[2:]
#                         ]
#                         table = insert_styled_table(doc, headers, rows)
#                         new_elements.append(table._element)
#                         continue

#                     # Section heading
#                     if line.startswith("**") or line.startswith("###"):
#                         header_text = line.strip("*# ").rstrip(":")
#                         new_para = doc.add_paragraph(header_text)
#                         new_para.style = "Table Column Heading"
#                         new_para.paragraph_format.space_after = Pt(4)
#                         new_elements.append(new_para._element)
#                         i += 1
#                         continue

         

#                     # Detect markdown bullets ("- " or "• ")
#                     if line.startswith("- ") or line.startswith("• "):
#                         bullet_text = line[2:].strip() if line.startswith("- ") else line[1:].strip()
#                         new_para = doc.add_paragraph(bullet_text, style="List Bullet 2")
#                         new_para.paragraph_format.left_indent = Pt(18)
#                         new_para.paragraph_format.space_after = Pt(2)
#                         new_elements.append(new_para._element)
#                         i += 1
#                         continue

#                     # Otherwise, add as regular paragraph
#                     new_para = doc.add_paragraph(line)
#                     new_elements.append(new_para._element)
#                     i += 1


#                 # ⚡️ Insert all new elements once
#                 for element in reversed(new_elements):
#                     parent.insert(idx, element)
#                 return



#     # Replace placeholders with sections
#     replace_placeholder(doc, "<<EXEC_SUMMARY>>", summary_text)
#     replace_placeholder(doc, "<<OBJECTIVE>>", objective_text)
#     replace_placeholder(doc, "<<SCOPE_TEXT>>", scope_text)
#     replace_placeholder(doc, "<<RESOURCE_SCHEDULE>>", resource_schedule_text)
#     replace_placeholder(doc, "<<COMMUNICATION_PLAN>>", communication_plan_text)

#     return doc



# def generate_exec_summary_and_objective(reference_text, condensed_rfp, num_interfaces=113):
#     """Generate both sections and return as separate strings."""
#     client = AzureOpenAI(
#         azure_endpoint=os.getenv("AZURE_OPENAI_FRFP_ENDPOINT"),
#         api_key=os.getenv("AZURE_OPENAI_FRFP_KEY"),
#         api_version=os.getenv("AZURE_OPENAI_FRFP_VERSION")
#     )

#     prompt = get_executive_summary_and_objective_prompt(reference_text, condensed_rfp, num_interfaces)

#     response = client.chat.completions.create(
#         model="Codetest",
#         temperature=0.3,
#         max_tokens=2000,
#         messages=[{"role": "user", "content": prompt}]
#     )

#     full_output = response.choices[0].message.content.strip()

#     # --- Split into Executive Summary and Objective ---
#     exec_match = re.search(r"\*\*?Executive Summary\*\*?\s*(.*?)\s*(?=\*\*?Objective\*\*?)", full_output, re.S | re.I)
#     obj_match = re.search(r"\*\*?Objective\*\*?\s*(.*)", full_output, re.S | re.I)

#     exec_text = exec_match.group(1).strip() if exec_match else full_output
#     obj_text = obj_match.group(1).strip() if obj_match else ""

#     return exec_text, obj_text

# def generate_scope_sections(reference_text, condensed_rfp, num_interfaces=None):
#     """Generate combined 'In Scope', 'Prerequisites', 'Assumptions', and 'Out of Scope' section as one block."""


#     client = AzureOpenAI(
#         azure_endpoint=os.getenv("AZURE_OPENAI_FRFP_ENDPOINT"),
#         api_key=os.getenv("AZURE_OPENAI_FRFP_KEY"),
#         api_version=os.getenv("AZURE_OPENAI_FRFP_VERSION")
#     )

#     prompt = get_scope_prereq_assumptions_prompt(reference_text, condensed_rfp, num_interfaces)

#     response = client.chat.completions.create(
#         model="Codetest",
#         temperature=0.3,
#         max_tokens=1200,
#         messages=[{"role": "user", "content": prompt}]
#     )

#     return response.choices[0].message.content.strip()

# def generate_resource_schedule_and_commercial(reference_text, condensed_rfp):
#     """Generate Resource Schedule and Commercials section."""
#     client = AzureOpenAI(
#         azure_endpoint=os.getenv("AZURE_OPENAI_FRFP_ENDPOINT"),
#         api_key=os.getenv("AZURE_OPENAI_FRFP_KEY"),
#         api_version=os.getenv("AZURE_OPENAI_FRFP_VERSION")
#     )

#     prompt = get_resource_schedule_and_commercial_prompt(reference_text, condensed_rfp)

#     response = client.chat.completions.create(
#         model="Codetest",
#         temperature=0.3,
#         max_tokens=2000,
#         messages=[{"role": "user", "content": prompt}]
#     )

#     return response.choices[0].message.content.strip()
# def generate_communication_plan(reference_text, condensed_rfp):
#     """Generate Communication Plan section."""
#     client = AzureOpenAI(
#         azure_endpoint=os.getenv("AZURE_OPENAI_FRFP_ENDPOINT"),
#         api_key=os.getenv("AZURE_OPENAI_FRFP_KEY"),
#         api_version=os.getenv("AZURE_OPENAI_FRFP_VERSION")
#     )

#     prompt = get_communication_plan_prompt(reference_text, condensed_rfp)
#     response = client.chat.completions.create(
#         model="Codetest",
#         temperature=0.3,
#         max_tokens=2500,
#         messages=[{"role": "user", "content": prompt}]
#     )

#     return response.choices[0].message.content.strip()



# # -------------------------------------------------------
# # 3. STREAMLIT UI (Improved UX)
# # -------------------------------------------------------

# st.markdown("### 🧾 Step 1: Upload RFP Document")

# uploaded_file = st.file_uploader("Upload RFP (PDF or DOCX)", type=["pdf", "docx"])
# progress = st.progress(0)

# if uploaded_file:
#     # STEP 1: Extract content
#     with st.spinner("🔍 Extracting RFP content..."):
#         rfp_text = extract_text(uploaded_file)
#         time.sleep(1.5)
#     st.toast("RFP content extracted successfully ✅")
#     progress.progress(15)

#     with st.expander("📄 Preview Extracted RFP Text"):
#         st.text_area("", rfp_text[:2500] + "..." if len(rfp_text) > 2500 else rfp_text, height=250)

#     num_interfaces = st.number_input(
#             "Enter number of interfaces (if known):",
#             min_value=0, max_value=10000, value=113, step=1,
#             help="Used to adjust migration scale dynamically."
#         )


#     # STEP 2: Build or load knowledge base
#     st.markdown("### 🧠 Step 2: Build Knowledge Base")
#     with st.spinner("📚 Building or loading knowledge base..."):
#         knowledge_db = build_knowledge_base()
#         time.sleep(1.5)
#     st.toast("Knowledge base ready ✅")
#     progress.progress(35)

#     with st.spinner("🔎 Retrieving relevant reference documents..."):
#         retriever = knowledge_db.as_retriever(search_kwargs={"k": 3})
#         ref_docs = retriever.invoke(rfp_text)
#         reference_text = "\n\n".join([d.page_content for d in ref_docs])
#     st.toast(f"{len(ref_docs)} relevant reference documents retrieved ✅")
#     progress.progress(45)

#     # STEP 3: Generate Proposal Sections
#     st.markdown("### ✍️ Step 3: Generating Response")
#     with st.expander("🧾 Step 3.1: Generate Executive Summary & Objective", expanded=True):
#         with st.spinner("✍️ Generating Executive Summary & Objective..."):
#             exec_summary, objective = generate_exec_summary_and_objective(reference_text, rfp_text, num_interfaces)
#         st.toast("Executive Summary & Objective generated ✅")
#         progress.progress(60)

#         tab1, tab2 = st.tabs(["Executive Summary", "Objective"])
#         with tab1:
#             st.markdown(exec_summary)
#         with tab2:
#             st.markdown(objective)

#     # STEP 4: Generate Scope & Assumptions
#     with st.expander("📘 Step 3.2: Scope, Assumptions & Out of Scope"):
#         with st.spinner("🧩 Generating scope and assumptions..."):
#             scope_text = generate_scope_sections(reference_text, rfp_text, num_interfaces)
#         st.toast("Scope and Assumptions generated ✅")
#         progress.progress(75)
#         st.markdown(scope_text)

#     # STEP 5: Resource Schedule & Commercials
#     with st.expander("📊 Step 3.3: Resource Schedule & Commercials"):
#         with st.spinner("📈 Generating resource schedule and commercials..."):
#             resource_schedule_text = generate_resource_schedule_and_commercial(reference_text, rfp_text)
#         st.toast("Resource Schedule & Commercials ready ✅")
#         progress.progress(85)

#         tabs = st.tabs(["Resource Schedule", "Commercials"])
#         with tabs[0]:
#             st.markdown(resource_schedule_text)
#         with tabs[1]:
#             st.markdown("*(Add commercial details if available.)*")

#     # STEP 6: Communication Plan
#     with st.expander("📊 Step 3.4: Generating Communication Plan..."):
#         with st.spinner("📢 Generating Communication Plan..."):
#             communication_plan_text = generate_communication_plan(reference_text, rfp_text)
#         st.toast("Communication Plan generated ✅")
#         progress.progress(95)

#     # STEP 7: Insert into Template & Download
#     st.markdown("### 📦 Step 4: Review & Download Final Proposal")
#     template_path = "Template/PIPO TO IS Response Template.docx"

#     if not os.path.exists(template_path):
#         st.error(f"Template not found at {template_path}")
#     else:
#         final_doc = insert_executive_summary_into_template(
#             template_path,
#             summary_text=exec_summary,
#             objective_text=objective,
#             scope_text=scope_text,
#             resource_schedule_text=resource_schedule_text,
#             communication_plan_text=communication_plan_text
#         )

#         buffer = BytesIO()
#         final_doc.save(buffer)
#         buffer.seek(0)
#         progress.progress(100)

#         st.download_button(
#             label="📥 Download Final Proposal (DOCX)",
#             data=buffer,
#             file_name=f"RFP_Response_{uploaded_file.name.split('.')[0]}.docx",
#             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         )

#         st.success("✅ Proposal generated successfully! Ready for download.")




import streamlit as st
import os
import time
import re
from io import BytesIO
from dotenv import load_dotenv
from PyPDF2 import PdfReader
import docx
from docx import Document
from openai import AzureOpenAI
from langchain_openai import AzureOpenAIEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document as LDocument
from docx.shared import Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
# Assuming Modules/prompts exists and is correctly imported
from Modules.prompts import (
    get_executive_summary_and_objective_prompt,
    get_scope_prereq_assumptions_prompt,
    get_resource_schedule_and_commercial_prompt,
    get_communication_plan_prompt
)


# -------------------------------------------------------
# 1. SETUP
# -------------------------------------------------------
load_dotenv()
KNOWLEDGE_FOLDER = "Knowledge_Repo"
PERSIST_DIR = "chroma_db"

st.set_page_config(page_title="RFP Proposal AI Generator", layout="wide")

# Custom CSS for Professional Look (Mimics AutoRFP style)
st.markdown("""
<style>
/* Primary Brand Color */
:root {
    --primary-blue: #1A75E0;
    --light-blue-bg: #EAF3FF;
}

/* Centering and large text for the main header */
.main-header {
    text-align: center;
    color: #000;
    font-size: 3em;
    font-weight: 800;
    padding-top: 20px;
    padding-bottom: 5px;
}
.highlight-text {
    color: var(--primary-blue);
}
.sub-tagline {
    text-align: center;
    color: #555;
    font-size: 1.1em;
    padding-bottom: 40px;
}

/* Style for the upload boxes (the two blocks requested) */
.upload-card {
    border: 1px solid #E0E0E0;
    border-radius: 12px;
    padding: 30px 20px;
    text-align: center;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
    transition: all 0.2s ease-in-out;
    background-color: #F9F9F9;
    height: 100%;
    margin-bottom: 20px;
}
.upload-card:hover {
    box-shadow: 0 8px 16px rgba(0, 0, 0, 0.1);
    border-color: var(--primary-blue);
}
.upload-header {
    color: var(--primary-blue);
    font-weight: 600;
    margin-bottom: 10px;
}

/* Streamlit component tweaks */
div.stButton > button {
    background-color: var(--primary-blue);
    color: white;
    border-radius: 8px;
    border: none;
    padding: 10px 20px;
    font-weight: bold;
    transition: background-color 0.2s;
}
div.stButton > button:hover {
    background-color: #145CB0;
}
/* Style status box titles for consistency */
.stStatus [data-testid="stStatusContainer"] > div:first-child > div:first-child {
    font-weight: 600;
    color: #333;
}
</style>
""", unsafe_allow_html=True)


# -------------------------------------------------------
# 2. UTILITIES
# -------------------------------------------------------

def extract_text(file):
    """Extract text from PDF or DOCX"""
    if file.name.endswith(".pdf"):
        reader = PdfReader(file)
        return "\n".join([p.extract_text() or "" for p in reader.pages])
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    return ""

def build_knowledge_base(folder=KNOWLEDGE_FOLDER, persist_dir=PERSIST_DIR):
    """Build or load persistent Chroma vector DB from Knowledge_Repo"""
    # ... (rest of build_knowledge_base remains the same)
    os.makedirs(folder, exist_ok=True)
    os.makedirs(persist_dir, exist_ok=True)

    embedding_model = AzureOpenAIEmbeddings(
        model="text-embedding-ada-002",
        azure_endpoint=os.getenv("AZURE_OPENAI_EMD_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_EMD_KEY"),
        api_version=os.getenv("AZURE_OPENAI_EMD_VERSION")
    )

    if os.listdir(persist_dir):
        return Chroma(
            embedding_function=embedding_model,
            persist_directory=persist_dir,
            collection_name="rfp_responses"
        )

    docs = []
    for f in os.listdir(folder):
        if f.endswith((".pdf", ".docx")):
            path = os.path.join(folder, f)
            text = extract_text(open(path, "rb"))
            if text.strip():
                docs.append(LDocument(page_content=text, metadata={"source": f}))

    if not docs:
        raise ValueError(f"No readable files found in {folder}")

    return Chroma.from_documents(
        documents=docs,
        embedding=embedding_model,
        persist_directory=persist_dir,
        collection_name="rfp_responses"
    )

def apply_bullet_to_para(paragraph, list_id='1'):
    """
    Applies a dot bullet style (list level 0) using its XML structure.
    Uses numId='1' which is often the default bullet style in templates.
    """
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    
    # Set the list level (0 is the main level)
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    
    # Set the list ID (Most default templates use ID '1' for the first bullet definition)
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), list_id)
    
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)


def insert_executive_summary_into_template(
    template_path,
    summary_text,
    objective_text=None,
    scope_text=None,
    resource_schedule_text=None,
    communication_plan_text=None,
):
    """
    Replace placeholders in the template:
    <<EXEC_SUMMARY>>, <<OBJECTIVE>>, <<SCOPE_TEXT>>, <<RESOURCE_SCHEDULE>>, <<COMMUNICATION_PLAN>>
    Now includes robust bullet point handling.
    """

    doc = Document(template_path)

    def set_cell_shading(cell, fill_color):
        """Add shading (background color) to a table cell."""
        tc_pr = cell._element.tcPr
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_color)
        tc_pr.append(shd)

    def set_table_border_white(table, cell_margin=150):
        """Set all table borders to white (for clean, minimal look)."""
        tbl = table._element
        tbl_pr = tbl.tblPr
        tbl_borders = OxmlElement("w:tblBorders")

        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border_el = OxmlElement(f"w:{border_name}")
            border_el.set(qn("w:val"), "single")
            border_el.set(qn("w:sz"), "4")  # thin border
            border_el.set(qn("w:space"), "0")
            border_el.set(qn("w:color"), "FFFFFF")  # white
            tbl_borders.append(border_el)

        tbl_pr.append(tbl_borders)


    def insert_styled_table(parent, headers, rows):
        """Create a table styled similar to RFP objective section."""
        table = parent.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = "Table Grid"
        table.autofit = True

        # Header row styling
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h.strip()
            set_cell_shading(hdr_cells[i], "008FD3")  # blue header
            for run in hdr_cells[i].paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Data rows
        for r, row_data in enumerate(rows):
            cells = table.rows[r + 1].cells
            for c, val in enumerate(row_data):
                cells[c].text = str(val).strip()
                set_cell_shading(cells[c], "E7EEF7")  # light gray row
                cells[c].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cells[c].paragraphs[0].alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if c == 0 else WD_ALIGN_PARAGRAPH.LEFT
                )

        # Set uniform width
        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(3)

        # Apply white borders
        set_table_border_white(table)

        return table

    def replace_placeholder(doc, placeholder, new_text):
        if not new_text:
            return

        for para in doc.paragraphs:
            if placeholder in "".join(run.text for run in para.runs):
                parent = para._element.getparent()
                idx = parent.index(para._element)
                parent.remove(para._element)

                lines = [line.strip() for line in new_text.split("\n") if line.strip()]
                new_elements = []  # collect to insert once

                i = 0
                while i < len(lines):
                    line = lines[i]

                    # Markdown-style table
                    if line.startswith("|") and "|" in line:
                        table_lines = []
                        while i < len(lines) and lines[i].startswith("|"):
                            table_lines.append(lines[i])
                            i += 1
                        headers = [h.strip("* ") for h in table_lines[0].strip("|").split("|")]
                        rows = [
                            [c.strip() for c in r.strip("|").split("|")]
                            for r in table_lines[2:]
                        ]
                        table = insert_styled_table(doc, headers, rows)
                        new_elements.append(table._element)
                        continue

                    # Section heading
                    if line.startswith("**") or line.startswith("###"):
                        header_text = line.strip("*# ").rstrip(":")
                        new_para = doc.add_paragraph(header_text)
                        new_para.style = "Table Column Heading"
                        new_para.paragraph_format.space_after = Pt(4)
                        new_elements.append(new_para._element)
                        i += 1
                        continue

                    # Markdown bullets (FIXED: Use apply_bullet_to_para for robustness)
                    if line.startswith("- ") or line.startswith("• "):
                        bullet_text = line[2:].strip() if line.startswith("- ") else line[1:].strip()
                        new_para = doc.add_paragraph(bullet_text, style="List Bullet 2")
                        new_para.paragraph_format.left_indent = Pt(18)
                        new_para.paragraph_format.space_after = Pt(2)
                        new_elements.append(new_para._element)
                        i += 1
                        continue

                    # Regular text
                    new_para = doc.add_paragraph(line)
                    new_elements.append(new_para._element)
                    i += 1

                # ⚡️ Insert all new elements once
                for element in reversed(new_elements):
                    parent.insert(idx, element)
                return

    # Replace placeholders with sections
    replace_placeholder(doc, "<<EXEC_SUMMARY>>", summary_text)
    replace_placeholder(doc, "<<OBJECTIVE>>", objective_text)
    replace_placeholder(doc, "<<SCOPE_TEXT>>", scope_text)
    replace_placeholder(doc, "<<RESOURCE_SCHEDULE>>", resource_schedule_text)
    replace_placeholder(doc, "<<COMMUNICATION_PLAN>>", communication_plan_text)

    return doc


def generate_exec_summary_and_objective(reference_text, condensed_rfp, num_interfaces=113):
    # ... (function body remains the same)
    client = AzureOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_FRFP_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_FRFP_KEY"),
        api_version=os.getenv("AZURE_OPENAI_FRFP_VERSION")
    )

    prompt = get_executive_summary_and_objective_prompt(reference_text, condensed_rfp, num_interfaces)

    response = client.chat.completions.create(
        model="Codetest",
        temperature=0.3,
        max_tokens=2000,
        messages=[{"role": "user", "content": prompt}]
    )

    full_output = response.choices[0].message.content.strip()

    # --- Split into Executive Summary and Objective ---
    exec_match = re.search(r"\*\*?Executive Summary\*\*?\s*(.*?)\s*(?=\*\*?Objective\*\*?)", full_output, re.S | re.I)
    obj_match = re.search(r"\*\*?Objective\*\*?\s*(.*)", full_output, re.S | re.I)

    exec_text = exec_match.group(1).strip() if exec_match else full_output
    obj_text = obj_match.group(1).strip() if obj_match else ""

    return exec_text, obj_text

def generate_scope_sections(reference_text, condensed_rfp, num_interfaces=None):
    # ... (function body remains the same)
    client = AzureOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_FRFP_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_FRFP_KEY"),
        api_version=os.getenv("AZURE_OPENAI_FRFP_VERSION")
    )

    prompt = get_scope_prereq_assumptions_prompt(reference_text, condensed_rfp, num_interfaces)

    response = client.chat.completions.create(
        model="Codetest",
        temperature=0.3,
        max_tokens=1200,
        messages=[{"role": "user", "content": prompt}]
    )

    return response.choices[0].message.content.strip()

def generate_resource_schedule_and_commercial(reference_text, condensed_rfp):
    # ... (function body remains the same)
    client = AzureOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_FRFP_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_FRFP_KEY"),
        api_version=os.getenv("AZURE_OPENAI_FRFP_VERSION")
    )

    prompt = get_resource_schedule_and_commercial_prompt(reference_text, condensed_rfp)

    response = client.chat.completions.create(
        model="Codetest",
        temperature=0.3,
        max_tokens=2000,
        messages=[{"role": "user", "content": prompt}]
    )

    return response.choices[0].message.content.strip()

def generate_communication_plan(reference_text, condensed_rfp):
    # ... (function body remains the same)
    client = AzureOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_FRFP_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_FRFP_KEY"),
        api_version=os.getenv("AZURE_OPENAI_FRFP_VERSION")
    )

    prompt = get_communication_plan_prompt(reference_text, condensed_rfp)
    response = client.chat.completions.create(
        model="Codetest",
        temperature=0.3,
        max_tokens=2500,
        messages=[{"role": "user", "content": prompt}]
    )

    return response.choices[0].message.content.strip()


# -------------------------------------------------------
# 3. STREAMLIT UI (Revamped Professional Look)
# -------------------------------------------------------
# Centered, large header
st.markdown("<div class='main-header'>Automate Your <span class='highlight-text'>Proposal Response</span></div>", unsafe_allow_html=True)
st.markdown("<p class='sub-tagline'>Respond to RFPs in minutes with AI-driven content generation.</p>", unsafe_allow_html=True)

# --- Step 1: Upload ---
st.markdown("## 📥 Step 1: Upload Your RFP Document")
col1, col2 = st.columns(2)

# Create two styled blocks for file upload
with col1:
    st.markdown("<div class='upload-card'><p class='upload-header'>Upload RFP (PDF)</p>", unsafe_allow_html=True)
    pdf_uploader = st.file_uploader(" ", type=["pdf"], key="pdf_up", help="Upload the PDF version of your Request for Proposal.", label_visibility="collapsed")
    st.markdown("</div>", unsafe_allow_html=True)

with col2:
    st.markdown("<div class='upload-card'><p class='upload-header'>Upload RFP (DOCX)</p>", unsafe_allow_html=True)
    docx_uploader = st.file_uploader(" ", type=["docx"], key="docx_up", help="Upload the DOCX version of your Request for Proposal.", label_visibility="collapsed")
    st.markdown("</div>", unsafe_allow_html=True)

# Consolidate uploaded file check
uploaded_file = pdf_uploader or docx_uploader

# --- Dynamic Input ---
st.markdown("---")
st.markdown("### ⚙️ Proposal Configuration")
# num_interfaces = st.number_input(
#     "Enter number of interfaces (e.g., for migration scope):",
#     min_value=0, max_value=10000, value=113, step=1,
#     help="This value is used to contextualize the migration scale in the proposal text.",
#     key="interface_input"
# )
# st.markdown("---")

# --- Generate Button ---
generate_clicked = st.button("🚀 Generate Proposal Response", type="primary")

# --- Conditional Logic ---
if generate_clicked:
    if not uploaded_file:
        st.error("⚠️ Please upload an RFP document (PDF or DOCX) before generating.")
    else:
        st.markdown("### ✍️ Step 2: Generating Your Proposal Response")
        with st.spinner("Analyzing RFP and preparing your AI-driven proposal response..."):
            # Call your main generation logic here
            # result = generate_proposal(uploaded_file, num_interfaces)

            with st.status("🚀 Generating Proposal Sections...", expanded=True) as status:
        
                # STEP 1: Extract content
                st.write("1/6 🔎 Extracting RFP content...")
                rfp_text = extract_text(uploaded_file)
                time.sleep(1)
                # --- 🔍 Auto-detect number of interfaces / integrations from RFP text ---
                # import re

                keywords = [
                    "interfaces?", "integration points?", "flows?", "connections?",
                    "touchpoints?", "ICOs?", "IFlows?", "mappings?", "adapters?"
                ]

                pattern = r'\b(\d{1,5})\s+(?:' + "|".join(keywords) + r')\b'
                matches = re.findall(pattern, rfp_text, flags=re.IGNORECASE)

                if matches:
                    num_interfaces = max(map(int, matches))  # take largest number found
                    st.info(f"📊 Detected approximately **{num_interfaces}** integrations in RFP.")
                else:
                    st.warning("⚠️ No explicit integration count detected — using default or manual input.")

                
                if len(rfp_text.strip()) < 100:
                    status.update(label="Extraction Failed", state="error", expanded=False)
                    st.error("Could not extract enough text from the document. Please check the file.")
                    st.stop()
                
                st.success("1/6 ✅ RFP content extracted!")
                status.update(label="🚀 Generating Proposal Sections... (20% Complete)", state="running")

                # STEP 2: Build or load knowledge base & Retrieve context
                st.write("2/6 📚 Loading knowledge base and retrieving reference documents...")
                knowledge_db = build_knowledge_base()
                retriever = knowledge_db.as_retriever(search_kwargs={"k": 3})
                ref_docs = retriever.invoke(rfp_text)
                reference_text = "\n\n".join([d.page_content for d in ref_docs])
                st.success(f"2/6 ✅ Retrieved {len(ref_docs)} relevant reference documents!")
                status.update(label="🚀 Generating Proposal Sections... (40% Complete)", state="running")


                
                # STEP 3: Generate Core Sections
                st.write("3/6 ✍️ Generating Executive Summary and Objective...")
                exec_summary, objective = generate_exec_summary_and_objective(reference_text, rfp_text, num_interfaces)
                st.success("3/6 ✅ Executive Summary & Objective generated.")
                status.update(label="🚀 Generating Proposal Sections... (60% Complete)", state="running")

                # STEP 4: Generate Scope Sections
                st.write("4/6 🧩 Generating Scope, Assumptions, and Prerequisites...")
                scope_text = generate_scope_sections(reference_text, rfp_text, num_interfaces)
                st.success("4/6 ✅ Scope and Assumptions section generated.")
                status.update(label="🚀 Generating Proposal Sections... (75% Complete)", state="running")
                
                # STEP 5: Resource Schedule & Commercials
                st.write("5/6 📊 Generating Resource Schedule and Commercials...")
                resource_schedule_text = generate_resource_schedule_and_commercial(reference_text, rfp_text)
                st.success("5/6 ✅ Resource Schedule and Commercials generated.")
                status.update(label="🚀 Generating Proposal Sections... (85% Complete)", state="running")

                # STEP 6: Communication Plan
                st.write("6/6 📢 Generating Communication Plan...")
                communication_plan_text = generate_communication_plan(reference_text, rfp_text)
                st.success("6/6 ✅ Communication Plan generated.")
                status.update(label="✅ Proposal Content Complete!", state="complete", expanded=False)

            
            # --- Proposal Preview ---
            st.markdown("## 🔍 Step 2: Review and Edit Content")
            st.info("Review the AI-generated sections below before downloading the final document.")
            
            tab1, tab2, tab3, tab4, tab5 = st.tabs([
                "Executive Summary", "Objective", "Scope & Assumptions", 
                "Resource & Schedule", "Communication Plan"
            ])
            
            with tab1: st.markdown(exec_summary)
            with tab2: st.markdown(objective)
            with tab3: st.markdown(scope_text)
            with tab4: st.markdown(resource_schedule_text)
            with tab5: st.markdown(communication_plan_text)
            
            # --- Download Section ---
            st.markdown("---")
            st.markdown("## 📦 Step 3: Final Document Generation & Download")
            template_path = "Template/PIPO TO IS Response Template.docx"

            if not os.path.exists(template_path):
                st.error(f"Template not found at {template_path}. Cannot generate final DOCX.")
            else:
                st.write("Compiling content into DOCX template...")
                final_doc = insert_executive_summary_into_template(
                    template_path,
                    summary_text=exec_summary,
                    objective_text=objective,
                    scope_text=scope_text,
                    resource_schedule_text=resource_schedule_text,
                    communication_plan_text=communication_plan_text
                )

                buffer = BytesIO()
                final_doc.save(buffer)
                buffer.seek(0)
                
                st.markdown("<br>", unsafe_allow_html=True)
                st.download_button(
                    label="🚀 Download Final RFP Proposal (DOCX)",
                    data=buffer,
                    file_name=f"RFP_Response_{uploaded_file.name.split('.')[0]}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )


            st.success("✅ Proposal response generated successfully!")
            # st.write(result)
else:
    st.info("Upload your RFP and enter configuration details, then click **Generate Proposal Response**.")





# # Centered, large header
# st.markdown("<div class='main-header'>Automate Your <span class='highlight-text'>Proposal Response</span></div>", unsafe_allow_html=True)
# st.markdown("<p class='sub-tagline'>Respond to RFPs in minutes with AI-driven content generation.</p>", unsafe_allow_html=True)


# # --- Step 1: Upload ---
# st.markdown("## 📥 Step 1: Upload Your RFP Document")
# col1, col2 = st.columns(2)

# # Create two styled blocks for file upload
# with col1:
#     st.markdown("<div class='upload-card'><p class='upload-header'>Upload RFP (PDF)</p>", unsafe_allow_html=True)
#     pdf_uploader = st.file_uploader(" ", type=["pdf"], key="pdf_up", help="Upload the PDF version of your Request for Proposal.", label_visibility="collapsed")
#     st.markdown("</div>", unsafe_allow_html=True)

# with col2:
#     st.markdown("<div class='upload-card'><p class='upload-header'>Upload RFP (DOCX)</p>", unsafe_allow_html=True)
#     docx_uploader = st.file_uploader(" ", type=["docx"], key="docx_up", help="Upload the DOCX version of your Request for Proposal.", label_visibility="collapsed")
#     st.markdown("</div>", unsafe_allow_html=True)

# # Consolidate the uploaded file check
# uploaded_file = pdf_uploader or docx_uploader
#         # --- Dynamic Input ---
# st.markdown("---")
# st.markdown("### ⚙️ Proposal Configuration")
# num_interfaces = st.number_input(
#             "Enter number of interfaces (e.g., for migration scope):",
#             min_value=0, max_value=10000, value=113, step=1,
#             help="This value is used to contextualize the migration scale in the proposal text.",
#             key="interface_input"
#         )
# st.markdown("---")

# if uploaded_file:
#     # Use st.status for a better, more detailed UX during processing
#     with st.status("🚀 Generating Proposal Sections...", expanded=True) as status:
        
#         # STEP 1: Extract content
#         st.write("1/6 🔎 Extracting RFP content...")
#         rfp_text = extract_text(uploaded_file)
#         time.sleep(1)
        
#         if len(rfp_text.strip()) < 100:
#              status.update(label="Extraction Failed", state="error", expanded=False)
#              st.error("Could not extract enough text from the document. Please check the file.")
#              st.stop()
        
#         st.success("1/6 ✅ RFP content extracted!")
#         status.update(label="🚀 Generating Proposal Sections... (20% Complete)", state="running")

#         # STEP 2: Build or load knowledge base & Retrieve context
#         st.write("2/6 📚 Loading knowledge base and retrieving reference documents...")
#         knowledge_db = build_knowledge_base()
#         retriever = knowledge_db.as_retriever(search_kwargs={"k": 3})
#         ref_docs = retriever.invoke(rfp_text)
#         reference_text = "\n\n".join([d.page_content for d in ref_docs])
#         st.success(f"2/6 ✅ Retrieved {len(ref_docs)} relevant reference documents!")
#         status.update(label="🚀 Generating Proposal Sections... (40% Complete)", state="running")


        
#         # STEP 3: Generate Core Sections
#         st.write("3/6 ✍️ Generating Executive Summary and Objective...")
#         exec_summary, objective = generate_exec_summary_and_objective(reference_text, rfp_text, num_interfaces)
#         st.success("3/6 ✅ Executive Summary & Objective generated.")
#         status.update(label="🚀 Generating Proposal Sections... (60% Complete)", state="running")

#         # STEP 4: Generate Scope Sections
#         st.write("4/6 🧩 Generating Scope, Assumptions, and Prerequisites...")
#         scope_text = generate_scope_sections(reference_text, rfp_text, num_interfaces)
#         st.success("4/6 ✅ Scope and Assumptions section generated.")
#         status.update(label="🚀 Generating Proposal Sections... (75% Complete)", state="running")
        
#         # STEP 5: Resource Schedule & Commercials
#         st.write("5/6 📊 Generating Resource Schedule and Commercials...")
#         resource_schedule_text = generate_resource_schedule_and_commercial(reference_text, rfp_text)
#         st.success("5/6 ✅ Resource Schedule and Commercials generated.")
#         status.update(label="🚀 Generating Proposal Sections... (85% Complete)", state="running")

#         # STEP 6: Communication Plan
#         st.write("6/6 📢 Generating Communication Plan...")
#         communication_plan_text = generate_communication_plan(reference_text, rfp_text)
#         st.success("6/6 ✅ Communication Plan generated.")
#         status.update(label="✅ Proposal Content Complete!", state="complete", expanded=False)

    
#     # --- Proposal Preview ---
#     st.markdown("## 🔍 Step 2: Review and Edit Content")
#     st.info("Review the AI-generated sections below before downloading the final document.")
    
#     tab1, tab2, tab3, tab4, tab5 = st.tabs([
#         "Executive Summary", "Objective", "Scope & Assumptions", 
#         "Resource & Schedule", "Communication Plan"
#     ])
    
#     with tab1: st.markdown(exec_summary)
#     with tab2: st.markdown(objective)
#     with tab3: st.markdown(scope_text)
#     with tab4: st.markdown(resource_schedule_text)
#     with tab5: st.markdown(communication_plan_text)
    
#     # --- Download Section ---
#     st.markdown("---")
#     st.markdown("## 📦 Step 3: Final Document Generation & Download")
#     template_path = "Template/PIPO TO IS Response Template.docx"

#     if not os.path.exists(template_path):
#         st.error(f"Template not found at {template_path}. Cannot generate final DOCX.")
#     else:
#         st.write("Compiling content into DOCX template...")
#         final_doc = insert_executive_summary_into_template(
#             template_path,
#             summary_text=exec_summary,
#             objective_text=objective,
#             scope_text=scope_text,
#             resource_schedule_text=resource_schedule_text,
#             communication_plan_text=communication_plan_text
#         )

#         buffer = BytesIO()
#         final_doc.save(buffer)
#         buffer.seek(0)
        
#         st.markdown("<br>", unsafe_allow_html=True)
#         st.download_button(
#             label="🚀 Download Final RFP Proposal (DOCX)",
#             data=buffer,
#             file_name=f"RFP_Response_{uploaded_file.name.split('.')[0]}.docx",
#             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         )
#         st.success("Proposal document ready for download!")
